"""Drift-guard: assert numerical claims in the manuscript DOCX match
``analysis_results.json``.

Run before release (v2.0-rev onwards; added in response to internal
peer-review Round 6 M1).

Usage
-----
    python src/render_docx_stub.py \
        --json analysis_results.json \
        --docx tables/tables_JCRCO_rev_v2.docx

Exit codes
----------
0   all tolerated values match
1   at least one mismatch detected (details printed to stderr)
2   an input file was missing

The guard is intentionally conservative: only values explicitly listed
in ``_CANONICAL_CLAIMS`` are checked. Extending the list requires a
companion update to the manuscript.
"""
from __future__ import annotations

import argparse
import json
import pathlib
import re
import sys


_CANONICAL_CLAIMS = {
    # (series_key, field)          : expected numeric OR (low, high) CI
    ("inc_asr_weiblich",  "APC")        : -0.44,
    ("inc_asr_männlich",  "APC")        : -0.37,
    ("mort_asr_weiblich", "APC")        : -0.54,
    ("mort_asr_männlich", "APC")        : -0.23,
    ("inc_asr_weiblich",  "MK_Z")       : -2.236,
    ("inc_asr_männlich",  "MK_Z")       : -2.518,
    ("mort_asr_weiblich", "MK_Z")       : -3.583,
    ("mort_asr_männlich", "MK_Z")       : -2.141,
    ("inc_asr_weiblich",  "durbin_watson"): 0.96,
    ("inc_asr_männlich",  "durbin_watson"): 0.79,
    ("mort_asr_weiblich", "durbin_watson"): 2.28,
    ("mort_asr_männlich", "durbin_watson"): 1.42,
}

# tolerance (absolute) for float comparison
ATOL = 0.01


def _load_docx_text(p: pathlib.Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        print(f"[render_docx_stub] python-docx not installed: {e}",
              file=sys.stderr)
        raise
    doc = Document(str(p))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def check(json_path: pathlib.Path, docx_path: pathlib.Path) -> int:
    if not json_path.exists():
        print(f"[render_docx_stub] JSON not found: {json_path}",
              file=sys.stderr)
        return 2
    if not docx_path.exists():
        print(f"[render_docx_stub] DOCX not found: {docx_path}",
              file=sys.stderr)
        return 2

    data = json.loads(json_path.read_text())
    docx_text = _load_docx_text(docx_path)

    mismatches: list[tuple[str, str, float, float | None]] = []
    for (series, field), expected in _CANONICAL_CLAIMS.items():
        series_data = data.get(series, {})
        actual = series_data.get(field)
        if actual is None:
            mismatches.append((series, field, float("nan"), None))
            continue
        if abs(float(actual) - float(expected)) > ATOL:
            mismatches.append((series, field, float(actual), float(expected)))
            continue
        # also check the DOCX text contains the numeric string to one decimal
        if isinstance(expected, float):
            candidates = {f"{expected:.3f}", f"{expected:.2f}",
                          f"{expected:+.3f}".lstrip("+"), f"{expected:+.2f}".lstrip("+")}
        else:
            candidates = {str(expected)}
        if not any(c in docx_text for c in candidates):
            mismatches.append((series, field + "/not in DOCX", float(actual),
                               float(expected)))

    if mismatches:
        print("[render_docx_stub] DRIFT DETECTED:", file=sys.stderr)
        for m in mismatches:
            print("  ", m, file=sys.stderr)
        return 1
    print("[render_docx_stub] OK — all canonical claims match.")
    return 0


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--json", required=True, type=pathlib.Path)
    ap.add_argument("--docx", required=True, type=pathlib.Path)
    args = ap.parse_args(argv)
    return check(args.json, args.docx)


if __name__ == "__main__":
    raise SystemExit(main())
